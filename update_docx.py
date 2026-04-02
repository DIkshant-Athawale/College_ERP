import os
from docx import Document

def update_document(doc_path, output_path):
    doc = Document(doc_path)

    for p in doc.paragraphs:
        if "organized into five tabs" in p.text:
            p.text = p.text.replace(
                "five tabs —Dashboard, Attendance, Assignments, Unit Tests, and Internal Marks",
                "eight tabs — Dashboard, Courses, Attendance, Assignments, Unit Tests, Internal Marks, Timetable, and Enrollment"
            )
        if "Seven tabs to switch" in p.text:
            p.text = p.text.replace("Seven tabs", "Eight tabs")
        
        # Add the 5.9 Enrollment Tab and rename 5.9 Notice Management to 5.10 Notice Management
        if "5.9 Notice Management" in p.text:
            p.text = p.text.replace("5.9 Notice Management", "5.10 Notice Management")
    
    # Insert "5.9 Enrollment Tab" right before the paragraph containing "5.10 Notice Management"
    for i, p in enumerate(doc.paragraphs):
        if "5.10 Notice Management" in p.text:
            # Insert the new heading before it
            new_p_heading = p.insert_paragraph_before("5.9 Enrollment Tab", style=p.style)
            new_p_content = new_p_heading.insert_paragraph_before(
                "This tab handles manual enrollments. Faculty select an assigned Professional or Open Elective course and view two lists: eligible students and currently enrolled students. Faculty can seamlessly enroll or drop students as needed, keeping the enrollment records up-to-date."
            )
            break

    # Also update tables for Student Data in 6.1
    for table in doc.tables:
        # Check if it's the 6.1 Data/Description table
        found_profile = False
        for row in table.rows:
            if row.cells[0].text.strip() == "Profile":
                found_profile = True
                break
        
        if found_profile:
            # We add a row for Current Status
            row_cells = table.add_row().cells
            row_cells[0].text = "Current Status"
            row_cells[1].text = "Student's current academic standing (e.g. Good Standing)"
            
            # We add a row for Essential Links
            row_cells = table.add_row().cells
            row_cells[0].text = "Essential Links"
            row_cells[1].text = "Admin-curated list of important institutional links"

    doc.save(output_path)
    print("Document successfully updated.")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_doc = os.path.join(script_dir, "thesis_formatted.docx")
    output_doc = os.path.join(script_dir, "thesis_formatted_updated.docx")
    update_document(input_doc, output_doc)
